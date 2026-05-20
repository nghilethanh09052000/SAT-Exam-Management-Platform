#!/usr/bin/env python3
"""
json_to_docx.py — Convert a raw SAT JSON file to an importable .docx file.

The output follows DOCX-TEMPLATE.md exactly so it can be uploaded directly to
the platform and parsed by the Mammoth.js + docx-parser pipeline.

Parseable fields (recognised by the platform parser):
  **Module N: Reading and Writing** / **Module N: Math**  ← bold heading
  **Question N**                                           ← bold heading
  - **Text:** [passage]
  - **Question:** [stem]
  - **Options:**
  - **A) correct text**                                    ← bold = correct
  - B) wrong text                                          ← plain = wrong
  - **Answer:** value1 | value2                            ← short answer

Extra metadata fields (plain text bullets — parser ignores them, visible for review):
  - difficulty: Easy
  - skill: Words in Context
  - explanation: ...

Usage
-----
    cd pipeline/sat-pipeline
    python3 scripts/json_to_docx.py output/raw/30116.json
    python3 scripts/json_to_docx.py output/raw/30116.json -o ~/Desktop/out.docx
    python3 scripts/json_to_docx.py output/raw/30116.json --section "Math"
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

from utils.classifier import classify_question

# ---------------------------------------------------------------------------
# Section name → DOCX module name
# Covers both the satgpt format and bluebooky format.
# ---------------------------------------------------------------------------

_SECTION_TO_MODULE: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"reading.*writing.*module\s*1", re.I), "Module 1: Reading and Writing"),
    (re.compile(r"reading.*writing.*module\s*2", re.I), "Module 2: Reading and Writing"),
    (re.compile(r"math.*module\s*1",              re.I), "Module 1: Math"),
    (re.compile(r"math.*module\s*2",              re.I), "Module 2: Math"),
    # Fallback: single-module files (bluebooky section = "rw" / "math")
    (re.compile(r"^(reading.*writing|rw)$",       re.I), "Module 1: Reading and Writing"),
    (re.compile(r"^math$",                         re.I), "Module 1: Math"),
]


def _module_name(section_name: str) -> str:
    for pattern, name in _SECTION_TO_MODULE:
        if pattern.search(section_name):
            return name
    return f"Module 1: {section_name}"


def _section_key(section_name: str) -> str:
    return "math" if "math" in section_name.lower() else "rw"


# ---------------------------------------------------------------------------
# HTML → plain text
# ---------------------------------------------------------------------------

_HTML_TAG_RE = re.compile(r"<[^>]+>")
_WHITESPACE_RE = re.compile(r"\s+")


def _plain(raw: str | None) -> str:
    if not raw:
        return ""
    text = BeautifulSoup(raw, "html.parser").get_text(separator=" ")
    return _WHITESPACE_RE.sub(" ", text).strip()


# ---------------------------------------------------------------------------
# python-docx helpers
# ---------------------------------------------------------------------------

def _bold_paragraph(doc: Document, text: str) -> None:
    """Add a standalone bold paragraph (module / question heading)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def _blank(doc: Document) -> None:
    doc.add_paragraph()


def _bullet_plain(doc: Document, text: str) -> None:
    """Bullet list item — entirely plain text."""
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def _bullet_field(doc: Document, field: str, content: str) -> None:
    """
    Bullet list item with a bold field name followed by plain content.
    Produces: <li><strong>Field:</strong> content</li>
    Mammoth maps this to the recognised field markers.
    """
    p = doc.add_paragraph(style="List Bullet")
    run_label = p.add_run(f"{field}: ")
    run_label.bold = True
    if content:
        p.add_run(content)


def _bullet_option(doc: Document, label: str, text: str, is_correct: bool) -> None:
    """
    Bullet list item for an A/B/C/D option.
    Correct answer → entire line bold (parser reads bold = correct).
    Wrong answer   → plain text.
    """
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{label}) {text}")
    run.bold = is_correct


# ---------------------------------------------------------------------------
# Core builder
# ---------------------------------------------------------------------------

def build_docx(data: dict, section_filter: str | None = None) -> Document:
    doc = Document()

    # Remove default empty paragraph that python-docx inserts
    for element in list(doc.element.body):
        doc.element.body.remove(element)

    first_section = True
    for section in data.get("sections", []):
        section_name = section.get("name", "")
        if section_filter and section_filter.lower() not in section_name.lower():
            continue

        module = _module_name(section_name)
        sk = _section_key(section_name)
        questions = section.get("questions", [])

        # Blank line before every module heading except the first
        if not first_section:
            _blank(doc)
        first_section = False

        # Module heading — bold paragraph
        _bold_paragraph(doc, module)

        for q_idx, q in enumerate(questions, 1):
            _blank(doc)

            # Question heading — bold paragraph, number from JSON or sequential
            q_number = q.get("number", q_idx)
            _bold_paragraph(doc, f"Question {q_number}")

            _blank(doc)

            # ── Extra metadata (plain bullets — parser ignores, human-readable) ──
            difficulty = (q.get("difficulty") or "").strip()
            if difficulty:
                _bullet_plain(doc, f"difficulty: {difficulty}")

            category = classify_question(
                text=_plain(q.get("content")),
                section=sk,  # type: ignore[arg-type]
                domain=q.get("domain"),
            )
            if category and category != "Uncategorized":
                _bullet_plain(doc, f"skill: {category}")

            _blank(doc)

            # ── Passage (Text field) ─────────────────────────────────────────
            passage = _plain(q.get("passage"))
            if passage:
                _bullet_field(doc, "Text", passage)
                _blank(doc)

            # ── Question stem ────────────────────────────────────────────────
            stem = _plain(q.get("content"))
            _bullet_field(doc, "Question", stem)

            _blank(doc)

            # ── Options or Answer ────────────────────────────────────────────
            correct = (q.get("correctAnswer") or "").strip().upper()
            options: dict[str, str] = q.get("options") or {}
            is_mc = any((options.get(lbl) or "").strip() for lbl in ["A", "B", "C", "D"])

            if is_mc:
                # Options section header (bold field name, no content)
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run("Options:")
                run.bold = True

                _blank(doc)

                for label in ["A", "B", "C", "D"]:
                    opt_text = _plain(options.get(label))
                    if not opt_text:
                        continue
                    is_correct = correct == label
                    _bullet_option(doc, label, opt_text, is_correct)
                    _blank(doc)
            else:
                # Short answer — Answer field with accepted values separated by " | "
                answer = (q.get("correctAnswer") or "").strip()
                p = doc.add_paragraph(style="List Bullet")
                run_label = p.add_run("Answer: ")
                run_label.bold = True
                run_val = p.add_run(answer)
                run_val.bold = True  # highlight the answer value too

            # ── Explanation (plain bullet — parser ignores, visible in file) ─
            explanation = _plain(
                q.get("explanation_text") or q.get("explanation")
            )
            if explanation:
                _blank(doc)
                _bullet_plain(doc, f"explanation: {explanation}")

    return doc


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert a raw SAT JSON file to an importable .docx (DOCX-TEMPLATE format)."
    )
    parser.add_argument(
        "input",
        help="Path to the raw JSON file (e.g. output/raw/30116.json)",
    )
    parser.add_argument(
        "--output", "-o",
        help="Output .docx path. Defaults to output/docx/<stem>.docx",
    )
    parser.add_argument(
        "--section",
        help='Only include sections whose name contains this string (e.g. "Math")',
    )
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    if args.output:
        output_path = Path(args.output)
    else:
        output_dir = input_path.parent.parent / "docx"
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"{input_path.stem}.docx"

    data = json.loads(input_path.read_text(encoding="utf-8"))
    title = data.get("title", input_path.stem)

    print(f"Building DOCX for: {title}")
    doc = build_docx(data, section_filter=args.section)
    doc.save(str(output_path))
    print(f"Done → {output_path}")


if __name__ == "__main__":
    main()
