#!/usr/bin/env python3
"""
Export crawler JSON into the mapped SAT import format.

The generated files follow .docs/DOCX-TEMPLATE.md / template_sat_no_module.docx:

    00000_Question_N(Source)_Category_SC
    00001_[stimulus]
    00002_[prompt]
    00003_[short answer only]
    00004_[explanation, optional]
    00005_[difficulty, optional]
    00006_
    [choices, one marked with T (True)]
    ==End==

Outputs:
    output/mapped/<source>/txt/<input>.txt
    output/mapped/<source>/docx/<input>.docx
    output/mapped/<source>/json/<input>.json
    output/mapped/report.json
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Iterable

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = ROOT_DIR / "output"
DEFAULT_MAPPED_DIR = OUTPUT_DIR / "mapped"

sys.path.insert(0, str(ROOT_DIR))
from utils.classifier import classify_question  # noqa: E402


@dataclass
class MappedQuestion:
    number: int
    source_code: str
    category: str
    stimulus: str
    prompt: str
    choices: list[str]
    correct_choice_index: int | None
    short_answers: list[str]
    explanation: str
    difficulty: str
    source_file: str
    source_question_id: str
    split_method: str


_SPACE_RE = re.compile(r"[ \t\r\f\v]+")
_MULTI_NEWLINE_RE = re.compile(r"\n{3,}")
_TAG_CLEAN_RE = re.compile(r"[^A-Za-z0-9 _&/()+:.-]+")
_ANSWER_MARK_RE = re.compile(r"\s*T\s*\(True\)\s*$", re.I)
_QUESTION_SPLIT_PATTERNS = [
    re.compile(r"\b(Which choice .+)$", re.I | re.S),
    re.compile(r"\b(Which finding .+)$", re.I | re.S),
    re.compile(r"\b(Which quotation .+)$", re.I | re.S),
    re.compile(r"\b(Which statement .+)$", re.I | re.S),
    re.compile(r"\b(Which choice best .+)$", re.I | re.S),
    re.compile(r"\b(As used in the text, .+)$", re.I | re.S),
    re.compile(r"\b(Based on the text[s]?, .+)$", re.I | re.S),
    re.compile(r"\b(According to the text, .+)$", re.I | re.S),
    re.compile(r"\b(What is .+)$", re.I | re.S),
    re.compile(r"\b(What value .+)$", re.I | re.S),
    re.compile(r"\b(How many .+)$", re.I | re.S),
]


def main() -> None:
    parser = argparse.ArgumentParser(description="Export crawler JSON to mapped SAT import files.")
    parser.add_argument("--source", choices=["all", "satpro", "satgpt", "bluebooky", "realprep", "realprep_tests"], default="all")
    parser.add_argument("--output-dir", default=str(DEFAULT_MAPPED_DIR))
    parser.add_argument("--limit", type=int, help="Limit files per source, useful for checking output.")
    parser.add_argument("--skip-existing", action="store_true")
    args = parser.parse_args()

    out_dir = Path(args.output_dir)
    sources = _source_paths(args.source)
    report: dict[str, Any] = {"output_dir": str(out_dir), "sources": {}}

    for source, raw_dir in sources.items():
        paths = sorted(raw_dir.glob("*.json")) if raw_dir.exists() else []
        if args.limit is not None:
            paths = paths[: args.limit]
        source_report = export_source(source, paths, out_dir, skip_existing=args.skip_existing)
        report["sources"][source] = source_report

    out_dir.mkdir(parents=True, exist_ok=True)
    report_path = out_dir / "report.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Report: {report_path}")


def _source_paths(source: str) -> dict[str, Path]:
    all_sources = {
        # crawl_satpro stores SATGPT-origin/API-compatible exam JSON here.
        "satpro": OUTPUT_DIR / "raw",
        # crawl_satgpt.py uses this directory, but it is currently empty in this workspace.
        "satgpt": OUTPUT_DIR / "satgpt" / "raw",
        "bluebooky": OUTPUT_DIR / "bluebooky" / "raw",
        "realprep": OUTPUT_DIR / "realprep" / "raw",
        "realprep_tests": OUTPUT_DIR / "realprep_tests" / "raw",
    }
    if source == "all":
        return all_sources
    return {source: all_sources[source]}


def export_source(source: str, paths: list[Path], out_dir: Path, skip_existing: bool) -> dict[str, Any]:
    docx_dir = out_dir / source / "docx"
    txt_dir = out_dir / source / "txt"
    json_dir = out_dir / source / "json"
    for directory in (docx_dir, txt_dir, json_dir):
        directory.mkdir(parents=True, exist_ok=True)

    stats = {
        "input_files": len(paths),
        "exported_files": 0,
        "questions": 0,
        "split_methods": {},
        "empty_prompt": 0,
        "empty_stimulus": 0,
        "multiple_choice": 0,
        "short_answer": 0,
        "errors": [],
    }

    for index, path in enumerate(paths, start=1):
        txt_path = txt_dir / f"{path.stem}.txt"
        docx_path = docx_dir / f"{path.stem}.docx"
        json_path = json_dir / f"{path.stem}.json"
        if skip_existing and txt_path.exists() and docx_path.exists() and json_path.exists():
            continue
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
            questions = normalize_file(source, path, data)
            if not questions:
                continue
            txt_path.write_text(render_mapped_text(questions), encoding="utf-8")
            write_docx(questions, docx_path)
            json_path.write_text(
                json.dumps([asdict(q) for q in questions], ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            stats["exported_files"] += 1
            stats["questions"] += len(questions)
            for q in questions:
                stats["split_methods"][q.split_method] = stats["split_methods"].get(q.split_method, 0) + 1
                stats["empty_prompt"] += int(not q.prompt)
                stats["empty_stimulus"] += int(not q.stimulus)
                stats["multiple_choice"] += int(bool(q.choices))
                stats["short_answer"] += int(bool(q.short_answers))
            print(f"[{source} {index:3}/{len(paths)}] {path.name} -> {len(questions)} questions")
        except Exception as exc:  # keep batch export moving
            stats["errors"].append({"file": str(path), "error": str(exc)})
            print(f"[{source} {index:3}/{len(paths)}] ERROR {path.name}: {exc}", file=sys.stderr)

    return stats


def normalize_file(source: str, path: Path, data: dict[str, Any]) -> list[MappedQuestion]:
    if source in {"satpro", "satgpt"}:
        return normalize_satpro(path, data, source)
    if source == "bluebooky":
        return normalize_bluebooky(path, data)
    if source in {"realprep", "realprep_tests"}:
        return normalize_realprep(path, data)
    return []


def normalize_satpro(path: Path, data: dict[str, Any], source: str) -> list[MappedQuestion]:
    title = str(data.get("title") or path.stem)
    source_code = _source_code(title, fallback=f"{source}_{data.get('id') or path.stem}")
    mapped: list[MappedQuestion] = []
    global_num = 0

    for section in data.get("sections") or []:
        section_name = str(section.get("name") or "")
        section_key = "math" if "math" in section_name.lower() else "rw"
        for q in section.get("questions") or []:
            global_num += 1
            number = _int(q.get("number"), global_num)
            stimulus = _plain(q.get("passage"))
            prompt = _plain(q.get("content"))
            choices, correct_index = _choices_from_dict(q.get("options"), q.get("correctAnswer"))
            answer = str(q.get("correctAnswer") or "").strip()
            short_answers = [] if choices else _split_answers(answer)
            category = q.get("skill") or q.get("domain") or classify_question(
                text=" ".join([stimulus, prompt]),
                section=section_key,
                domain=q.get("domain"),
            )
            mapped.append(
                MappedQuestion(
                    number=number,
                    source_code=source_code,
                    category=_clean_category(category or section_name or "SAT"),
                    stimulus=stimulus,
                    prompt=prompt,
                    choices=choices,
                    correct_choice_index=correct_index,
                    short_answers=short_answers,
                    explanation=_plain(q.get("explanation_text") or q.get("explanation")),
                    difficulty=_clean_difficulty(q.get("difficulty")),
                    source_file=str(path.relative_to(ROOT_DIR)),
                    source_question_id=str(q.get("id") or ""),
                    split_method="native_fields",
                )
            )
    return mapped


def normalize_bluebooky(path: Path, data: dict[str, Any]) -> list[MappedQuestion]:
    test_date = str(data.get("test_date") or path.stem)
    section = str(data.get("section") or "")
    section_key = "math" if section == "math" else "rw"
    source_code = _source_code(f"bluebooky_{test_date}_{section}", fallback=path.stem)
    mapped: list[MappedQuestion] = []

    for index, q in enumerate(data.get("questions") or [], start=1):
        stimulus_parts = [_bluebooky_text(q.get("passage"))]
        viz_text = _viz_to_text(q.get("viz_data"))
        if viz_text:
            stimulus_parts.append(viz_text)
        if q.get("image_url"):
            stimulus_parts.append(f"Image: {q.get('image_url')}")
        choices = [_plain(opt.get("text")) for opt in q.get("options") or [] if isinstance(opt, dict)]
        correct = str(q.get("correct_answer") or "").strip().upper()
        correct_index = "ABCDEFGHIJKLMNOPQRSTUVWXYZ".find(correct) if correct else None
        if correct_index is not None and correct_index < 0:
            correct_index = None
        mapped.append(
            MappedQuestion(
                number=index,
                source_code=source_code,
                category=_clean_category(classify_question(_bluebooky_text(q.get("prompt")), section_key, None) or section or "SAT"),
                stimulus="\n\n".join(part for part in stimulus_parts if part),
                prompt=_bluebooky_text(q.get("prompt")),
                choices=choices,
                correct_choice_index=correct_index,
                short_answers=[] if choices else _split_answers(q.get("correct_answer")),
                explanation="",
                difficulty="",
                source_file=str(path.relative_to(ROOT_DIR)),
                source_question_id=str(q.get("id") or ""),
                split_method="native_fields",
            )
        )
    return mapped


def normalize_realprep(path: Path, data: dict[str, Any]) -> list[MappedQuestion]:
    title = " ".join(
        part
        for part in [data.get("course_title"), data.get("title") or data.get("list_title")]
        if part
    ) or path.stem
    source_code = _source_code(title, fallback=f"{data.get('source') or 'realprep'}_{data.get('quiz_id') or path.stem}")
    section_key = "math" if "math" in str(data.get("section") or data.get("course_title") or "").lower() else "rw"
    mapped: list[MappedQuestion] = []

    for index, q in enumerate(data.get("questions") or [], start=1):
        number = _int(q.get("number"), index)
        stimulus, prompt, split_method = split_realprep_question(q)
        choices, correct_index = _choices_from_realprep(q.get("choices"), q.get("correct_answer"))
        answer = q.get("correct_answer")
        short_answers = [] if choices else _split_answers(answer)
        raw_category = q.get("category") or data.get("course_title") or data.get("title")
        category = raw_category or classify_question(" ".join([stimulus, prompt]), section_key, None)
        mapped.append(
            MappedQuestion(
                number=number,
                source_code=source_code,
                category=_clean_category(category or ("Math" if section_key == "math" else "Reading and Writing")),
                stimulus=stimulus,
                prompt=prompt,
                choices=choices,
                correct_choice_index=correct_index,
                short_answers=short_answers,
                explanation=_plain(q.get("explanation_html")),
                difficulty="",
                source_file=str(path.relative_to(ROOT_DIR)),
                source_question_id=str(q.get("question_post_id") or q.get("question_pro_id") or ""),
                split_method=split_method,
            )
        )
    return mapped


def split_realprep_question(question: dict[str, Any]) -> tuple[str, str, str]:
    html_text = question.get("question_html") or ""
    soup = BeautifulSoup(html_text, "html.parser")
    split = soup.select_one("hr.lsqb-split")
    if split:
        before = _collect_sibling_html(split, previous=True)
        after = _collect_sibling_html(split, previous=False)
        return _repair_realprep_spacing(_plain(before)), _repair_realprep_spacing(_plain(after)), "realprep_lsqb_split"

    passage = soup.select_one(".lsqb-passage-content")
    if passage:
        prompt_nodes = _siblings_after(passage)
        prompt_html = "".join(str(node) for node in prompt_nodes)
        prompt = _repair_realprep_spacing(_plain(prompt_html))
        if prompt:
            return _repair_realprep_spacing(_plain(str(passage))), prompt, "realprep_passage_class"

    text = _repair_realprep_spacing(_plain(html_text) or _plain(question.get("question_text")))
    inferred_stimulus, inferred_prompt = _split_text_by_prompt_pattern(text)
    if inferred_stimulus and inferred_prompt:
        return inferred_stimulus, inferred_prompt, "realprep_prompt_pattern"

    # Math questions often consist only of one direct ask. In the mapped format,
    # keep those in 00002_ and leave 00001_ blank.
    if _looks_like_direct_question(text) or _is_math_question(question):
        return "", text, "direct_prompt_only"
    return text, "", "stimulus_only"


def _repair_realprep_spacing(text: str) -> str:
    if not text:
        return ""
    # Some RealPrep HTML has poem lines joined with no space/newline, e.g.
    # "wifeOf", "lifeAnd", "chief.I", "heIs". Restore readable line breaks.
    text = re.sub(r"(?<=[a-z])(?=Of\b)", "\n", text)
    text = re.sub(r"(?<=[a-z])(?=And\b)", "\n", text)
    text = re.sub(r"(?<=[a-z])(?=Is\b)", "\n", text)
    text = re.sub(r"(?<=\.)I(?=\s+am\b)", "\nI", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def render_mapped_text(questions: Iterable[MappedQuestion]) -> str:
    blocks = []
    for q in questions:
        lines = [
            f"00000_Question_{q.number}({q.source_code})_{q.category}{'_SC' if q.choices else ''}",
            f"00001_{q.stimulus}",
            f"00002_{q.prompt}",
        ]
        if q.short_answers:
            lines.append(f"00003_{' | '.join(q.short_answers)}")
        if q.explanation:
            lines.append(f"00004_{q.explanation}")
        if q.difficulty:
            lines.append(f"00005_{q.difficulty}")
        if q.choices:
            lines.append("00006_")
            for idx, choice in enumerate(q.choices):
                suffix = " T (True)" if q.correct_choice_index == idx else ""
                lines.append(f"{_ANSWER_MARK_RE.sub('', choice).strip()}{suffix}")
        lines.append("==End==")
        blocks.append("\n".join(lines))
    return "\n\n".join(blocks) + "\n"


def write_docx(questions: list[MappedQuestion], output_path: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    # Remove only the default empty paragraph. Keep sectPr so the DOCX has
    # valid page geometry for Word/LibreOffice renderers.
    body = doc.element.body
    for element in list(body):
        if element.tag.endswith("}p"):
            body.remove(element)

    for q_index, q in enumerate(questions):
        if q_index:
            doc.add_paragraph()
        block = render_mapped_text([q]).rstrip("\n").splitlines()
        for line_index, line in enumerate(block):
            p = doc.add_paragraph()
            _add_markdown_runs(p, line, force_bold=line_index == 0)
    doc.save(str(output_path))


def _plain(raw: Any) -> str:
    if raw is None:
        return ""
    text = str(raw).replace("\\n", "\n")
    if "<" in text and ">" in text:
        text = _html_to_markdown_text(text)
    text = text.replace("\u200b", "")
    text = _SPACE_RE.sub(" ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = _MULTI_NEWLINE_RE.sub("\n\n", text)
    return text.strip()


def _bluebooky_text(raw: Any) -> str:
    text = _plain(raw)
    text = re.sub(r"\+\+(.*?)\+\+", r"\1", text)
    return text.strip()


def _html_to_markdown_text(raw_html: str) -> str:
    soup = BeautifulSoup(raw_html, "html.parser")
    for noise in soup.select("script, style"):
        noise.decompose()
    return _node_to_markdown(soup).strip()


def _node_to_markdown(node: Any) -> str:
    if isinstance(node, NavigableString):
        return str(node)
    if not isinstance(node, Tag):
        return ""

    name = node.name.lower() if node.name else ""
    if name == "br":
        return "\n"
    if name == "hr":
        return "\n"

    inner = "".join(_node_to_markdown(child) for child in node.children)
    if name in {"em", "i"} and inner.strip():
        return f"*{inner.strip()}*"
    if name in {"strong", "b"} and inner.strip():
        return f"**{inner.strip()}**"
    if name == "u" and inner.strip():
        return inner.strip()
    if name in {"p", "div", "li", "tr", "table", "blockquote", "h1", "h2", "h3", "h4", "h5", "h6"}:
        return f"\n{inner.strip()}\n" if inner.strip() else "\n"
    return inner


def _add_markdown_runs(paragraph: Any, text: str, force_bold: bool = False) -> None:
    for chunk, bold, italic in _markdown_chunks(text):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.bold = force_bold or bold
        run.italic = italic


def _markdown_chunks(text: str) -> list[tuple[str, bool, bool]]:
    chunks: list[tuple[str, bool, bool]] = []
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end != -1:
                chunks.append((text[i + 2 : end], True, False))
                i = end + 2
                continue
            chunks.append(("**", False, False))
            i += 2
            continue
        if text[i] == "*":
            end = text.find("*", i + 1)
            if end != -1:
                chunks.append((text[i + 1 : end], False, True))
                i = end + 1
                continue
            chunks.append(("*", False, False))
            i += 1
            continue
        next_bold = text.find("**", i)
        next_italic = text.find("*", i)
        candidates = [pos for pos in (next_bold, next_italic) if pos != -1]
        next_marker = min(candidates) if candidates else len(text)
        if next_marker == i:
            chunks.append((text[i], False, False))
            i += 1
            continue
        chunks.append((text[i:next_marker], False, False))
        i = next_marker
    return chunks


def _collect_sibling_html(node: Tag, previous: bool) -> str:
    siblings = node.previous_siblings if previous else node.next_siblings
    parts = [str(sibling) for sibling in siblings]
    if previous:
        parts.reverse()
    return "".join(parts)


def _siblings_after(node: Tag) -> list[Any]:
    return [sibling for sibling in node.next_siblings]


def _split_text_by_prompt_pattern(text: str) -> tuple[str, str]:
    for pattern in _QUESTION_SPLIT_PATTERNS:
        match = pattern.search(text)
        if not match or match.start(1) <= 0:
            continue
        return text[: match.start(1)].strip(), match.group(1).strip()
    return "", ""


def _looks_like_direct_question(text: str) -> bool:
    if not text:
        return False
    return text.endswith("?") or bool(re.search(r"\b(which|what|how many|what value|if)\b", text[:80], re.I))


def _is_math_question(question: dict[str, Any]) -> bool:
    haystack = " ".join(str(question.get(key) or "") for key in ("category", "question_text", "question_html"))
    return bool(re.search(r"\b(math|equation|function|triangle|circle|linear|quadratic|percent|mean|median|slope)\b", haystack, re.I))


def _choices_from_dict(options: Any, correct_answer: Any) -> tuple[list[str], int | None]:
    if not isinstance(options, dict):
        return [], None
    labels = ["A", "B", "C", "D"]
    choices = [_plain(options.get(label)) for label in labels if _plain(options.get(label))]
    correct = str(correct_answer or "").strip().upper()
    correct_index = labels.index(correct) if correct in labels[: len(choices)] else None
    return choices, correct_index


def _choices_from_realprep(choices_data: Any, correct_answer: Any) -> tuple[list[str], int | None]:
    if not isinstance(choices_data, list):
        return [], None
    choices = [_plain(choice.get("text") or choice.get("html")) for choice in choices_data if isinstance(choice, dict)]
    correct_index = None
    for idx, choice in enumerate(choices_data):
        if isinstance(choice, dict) and choice.get("is_correct"):
            correct_index = idx
            break
    if correct_index is None:
        correct = str(correct_answer or "").strip().upper()
        labels = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
        if correct in labels[: len(choices)]:
            correct_index = labels.index(correct)
    return choices, correct_index


def _split_answers(answer: Any) -> list[str]:
    if answer is None:
        return []
    text = str(answer).strip()
    if not text:
        return []
    return [part.strip() for part in re.split(r"\s*(?:\||/|,|;)\s*", text) if part.strip()]


def _source_code(value: str, fallback: str) -> str:
    value = value or fallback
    value = re.sub(r"[^A-Za-z0-9]+", "", value)
    return value[:48] or re.sub(r"[^A-Za-z0-9]+", "", fallback)[:48] or "SAT"


def _clean_category(value: Any) -> str:
    text = _plain(value) or "SAT"
    text = text.replace("[", "(").replace("]", ")")
    text = _TAG_CLEAN_RE.sub("", text)
    text = re.sub(r"\s+", " ", text).strip(" _-")
    return text[:90] or "SAT"


def _clean_difficulty(value: Any) -> str:
    text = str(value or "").strip().title()
    return text if text in {"Easy", "Medium", "Hard"} else ""


def _int(value: Any, fallback: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return fallback


def _viz_to_text(viz: Any) -> str:
    if not isinstance(viz, dict):
        return ""
    if viz.get("type") == "table":
        title = _plain(viz.get("title"))
        headers = [str(h) for h in viz.get("headers") or []]
        rows = viz.get("rows") or []
        lines = [title] if title else []
        if headers:
            lines.append(" | ".join(headers))
        for row in rows:
            lines.append(" | ".join(str(cell) for cell in row))
        return "\n".join(lines).strip()
    return f"Visual data: {json.dumps(viz, ensure_ascii=False)}"


if __name__ == "__main__":
    main()
