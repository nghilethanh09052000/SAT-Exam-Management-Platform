import io
import tempfile
import zipfile
from collections import defaultdict
from datetime import datetime

import structlog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from temporalio import activity

from models import ExportResult, RawQuestion

log = structlog.get_logger()

_SECTION_FILENAME = {
    "Math": "Math.docx",
    "Reading & Writing": "Reading_and_Writing.docx",
}


def _build_docx(section: str, questions: list[RawQuestion]) -> bytes:
    """Build a DOCX for one section, questions grouped by domain."""
    doc = Document()

    # Title
    title = doc.add_heading(f"SAT Practice Questions — {section}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Generated {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    # Group by domain
    by_domain: dict[str, list[RawQuestion]] = defaultdict(list)
    for q in questions:
        by_domain[q.domain].append(q)

    for domain, domain_qs in sorted(by_domain.items()):
        doc.add_heading(domain, level=1)

        for i, q in enumerate(domain_qs, start=1):
            # Question header line
            header = doc.add_paragraph()
            run = header.add_run(f"Question {i}")
            run.bold = True
            run.font.size = Pt(11)
            if q.difficulty:
                diff_run = header.add_run(f"  [{q.difficulty}]")
                diff_run.font.size = Pt(9)
                diff_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # Question text
            doc.add_paragraph(q.question_text)

            # Choices
            for label, text in sorted(q.choices.items()):
                choice_para = doc.add_paragraph(style="List Bullet")
                choice_run = choice_para.add_run(f"{label})  ")
                choice_run.bold = True
                choice_para.add_run(text)

            # Answer + explanation
            answer_para = doc.add_paragraph()
            answer_run = answer_para.add_run(f"Answer: {q.correct_answer}")
            answer_run.bold = True
            answer_run.font.color.rgb = RGBColor(0x1A, 0x7A, 0x1A)

            if q.explanation:
                exp_para = doc.add_paragraph()
                exp_para.add_run("Explanation: ").bold = True
                exp_para.add_run(q.explanation)

            doc.add_paragraph("─" * 60)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@activity.defn
async def generate_and_upload_export(questions: list[RawQuestion]) -> ExportResult:
    """Generate one DOCX per section, zip them, and return the local ZIP path."""
    activity.heartbeat("generating DOCX files")

    # Split by section
    by_section: dict[str, list[RawQuestion]] = defaultdict(list)
    for q in questions:
        by_section[q.section].append(q)

    with tempfile.NamedTemporaryFile(suffix=".zip", delete=False) as tmp:
        tmp_path = tmp.name

    with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for section, section_qs in by_section.items():
            activity.heartbeat(f"building DOCX for {section}")
            docx_bytes = _build_docx(section, section_qs)
            filename = _SECTION_FILENAME.get(section, f"{section.replace(' ', '_')}.docx")
            zf.writestr(filename, docx_bytes)

    log.info(
        "export generated",
        zip_path=tmp_path,
        sections=list(by_section.keys()),
        questions=len(questions),
    )

    return ExportResult(
        zip_path=tmp_path,
        sections_exported=list(by_section.keys()),
        questions_exported=len(questions),
    )
