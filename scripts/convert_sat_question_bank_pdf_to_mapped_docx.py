#!/usr/bin/env python3
"""Convert a SAT Suite Question Bank PDF into the mapped DOCX import format.

The Question Bank PDFs often encode equations as positioned glyphs/images, so
plain text extraction loses important math. This converter therefore embeds a
cropped screenshot of each question area and also extracts answer choices /
accepted answers for grading.
"""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import tempfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from PIL import Image


PAGE_XML_WIDTH = 892
PAGE_XML_HEIGHT = 1262
QUESTION_TOP = 180


def run(cmd: list[str]) -> None:
    subprocess.run(cmd, check=True)


def read_pages(pdf_path: Path) -> list[str]:
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
        text_path = Path(tmp.name)
    try:
        run(["pdftotext", "-raw", str(pdf_path), str(text_path)])
        text = text_path.read_text(encoding="utf-8", errors="replace")
    finally:
        text_path.unlink(missing_ok=True)
    return text.split("\f")


def html_unescape_minimal(value: str) -> str:
    return (
        value.replace("&amp;", "&")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&quot;", '"')
        .replace("&#39;", "'")
    )


def extract_answer_tops(pdf_path: Path, work_dir: Path) -> dict[int, int]:
    xml_base = work_dir / "pdfxml"
    run(["pdftohtml", "-xml", "-q", str(pdf_path), str(xml_base)])
    xml_path = xml_base.with_suffix(".xml")
    root = ET.parse(xml_path).getroot()
    answer_tops: dict[int, int] = {}

    for page in root.findall("page"):
        page_number = int(page.attrib["number"])
        candidates: list[int] = []
        for text in page.findall("text"):
            top = int(float(text.attrib.get("top", "0")))
            content = "".join(text.itertext())
            content = html_unescape_minimal(content)
            if " Answer" in content or "Correct Answer:" in content:
                candidates.append(top)
        if candidates:
            answer_tops[page_number] = min(candidates)

    return answer_tops


def render_pages(pdf_path: Path, image_dir: Path, dpi: int) -> list[Path]:
    prefix = image_dir / "page"
    run(["pdftoppm", "-jpeg", "-r", str(dpi), str(pdf_path), str(prefix)])
    return sorted(image_dir.glob("page-*.jpg"))


def clean_line(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\ufeff", "").replace("\u00a0", " ")).strip()


def safe_header_category(value: str) -> str:
    cleaned = clean_line(value)
    cleaned = re.sub(r"[()]", "", cleaned)
    cleaned = cleaned.replace("_", " ")
    return cleaned or "Math"


def question_id(page_text: str, fallback: int) -> str:
    match = re.search(r"Question ID\s+([A-Za-z0-9]+)", page_text)
    return match.group(1) if match else f"page{fallback:03d}"


def split_question_answer(page_text: str, qid: str) -> tuple[str, str]:
    marker = re.search(rf"ID:\s*{re.escape(qid)}\s+Answer", page_text)
    if marker:
        return page_text[: marker.start()], page_text[marker.end() :]
    generic = re.search(r"\bCorrect Answer:\b", page_text)
    if generic:
        return page_text[: generic.start()], page_text[generic.start() :]
    return page_text, ""


def extract_options(question_text: str) -> list[tuple[str, str]]:
    lines = [clean_line(line) for line in question_text.splitlines()]
    options: list[tuple[str, str]] = []
    for idx, line in enumerate(lines):
        match = re.match(r"^([A-D])\.\s*(.*)$", line)
        if not match:
            continue
        label = match.group(1)
        content = match.group(2).strip()
        # Some formula-only choices extract as a blank label. Keep a stable,
        # nonempty placeholder; the embedded image preserves the original.
        if not content:
            content = f"Choice {label} (see image)"
        options.append((label, content))
    return options[:4]


def extract_correct_answer(answer_text: str) -> str | None:
    after_label = re.search(
        r"Correct Answer:\s*\n\s*([^\n]+?)(?=\n\s*(?:Rationale|Question Difficulty:)|$)",
        answer_text,
        re.S,
    )
    if after_label:
        value = clean_line(after_label.group(1))
        if value and value.lower() != "rationale":
            return value

    sentence = re.search(r"The correct answer is\s+(.+?)(?:\.|\n)", answer_text, re.I | re.S)
    if sentence:
        value = clean_line(sentence.group(1))
        if value:
            return value
    return None


def accepted_answers(answer_text: str, correct: str | None) -> list[str]:
    values: list[str] = []
    if correct and not re.fullmatch(r"[A-D]", correct.strip(), re.I):
        for part in re.split(r"\s*,\s*|\s+\|\s+", correct):
            part = re.sub(r"^(?:and|or)\s+", "", clean_line(part), flags=re.I)
            if part:
                values.append(part)

    examples = re.search(r"Note that\s+(.+?)\s+are examples of ways to enter a correct answer", answer_text, re.I | re.S)
    if examples:
        for part in re.split(r"\s*,\s*|\s+and\s+", clean_line(examples.group(1))):
            part = re.sub(r"^(?:and|or)\s+", "", part.strip(), flags=re.I)
            if part:
                values.append(part)

    deduped: list[str] = []
    seen = set()
    for value in values:
        if value not in seen:
            seen.add(value)
            deduped.append(value)
    return deduped


def extract_difficulty(answer_text: str) -> str | None:
    match = re.search(r"Question Difficulty:\s*\n?\s*(Easy|Medium|Hard)\b", answer_text, re.I)
    if not match:
        return None
    return match.group(1).capitalize()


def extract_pdf_metadata_value(answer_text: str, field: str, stop_fields: list[str]) -> str | None:
    lines = [clean_line(line) for line in answer_text.splitlines()]
    for idx, line in enumerate(lines):
        if line.lower() != field.lower():
            continue
        values: list[str] = []
        for next_line in lines[idx + 1:]:
            if not next_line:
                continue
            if any(next_line.lower() == stop.lower() for stop in stop_fields):
                break
            values.append(next_line)
        value = clean_line(" ".join(values))
        return value or None
    return None


def extract_category(answer_text: str) -> str:
    skill = extract_pdf_metadata_value(answer_text, "Skill", ["Difficulty", "Question Difficulty:", "Assessment", "Test", "Domain"])
    if skill:
        return safe_header_category(skill)

    domain = extract_pdf_metadata_value(answer_text, "Domain", ["Skill", "Difficulty", "Question Difficulty:", "Assessment", "Test"])
    if domain:
        return safe_header_category(domain)

    return "Math"


def infer_default_difficulty(pdf_path: Path) -> str | None:
    name = pdf_path.name.lower()
    if "easy" in name:
        return "Easy"
    if "medium" in name:
        return "Medium"
    if "hard" in name:
        return "Hard"
    return None


def latex_fraction(value: str) -> str:
    return re.sub(
        r"(?<![\w/.-])(-?\d+)\s*/\s*(\d+)(?![\w/.-])",
        lambda m: rf"\(\frac{{{m.group(1)}}}{{{m.group(2)}}}\)",
        value,
    )


def clean_rationale_lines(lines: list[str]) -> list[str]:
    cleaned: list[str] = []
    i = 0
    while i < len(lines):
        line = clean_line(lines[i])
        if not line:
            i += 1
            continue

        if line == "-" and i + 2 < len(lines):
            numerator = clean_line(lines[i + 1])
            denominator = clean_line(lines[i + 2])
            if re.fullmatch(r"\d+", numerator) and re.fullmatch(r"\d+", denominator):
                cleaned.append(f"-{numerator}/{denominator}")
                i += 3
                continue

        if re.fullmatch(r"-?\d+", line) and i + 1 < len(lines):
            denominator = clean_line(lines[i + 1])
            if re.fullmatch(r"\d+", denominator):
                cleaned.append(f"{line}/{denominator}")
                i += 2
                continue

        cleaned.append(line)
        i += 1

    return cleaned


def extract_rationale(answer_text: str) -> str | None:
    match = re.search(
        r"\bRationale\s*\n([\s\S]*?)(?=\n\s*Question Difficulty:|\n\s*Assessment\b|\Z)",
        answer_text,
        re.I,
    )
    if not match:
        return None

    lines = match.group(1).splitlines()
    rationale = clean_line(" ".join(clean_rationale_lines(lines)))
    if not rationale:
        return None

    # Keep common math fractions import-friendly when PDF text extraction splits
    # stacked fractions into separate lines.
    return latex_fraction(rationale)


def extract_prompt(question_text: str) -> str:
    lines = [clean_line(line) for line in question_text.splitlines()]
    body_lines: list[str] = []
    for line in lines:
        if not line:
            continue
        if line.startswith("Question ID") or re.match(r"^ID:\s*[A-Za-z0-9]+$", line):
            continue
        if re.match(r"^[A-D]\.", line):
            break
        body_lines.append(line)

    body = clean_line(" ".join(body_lines))
    question_sentences = re.findall(r"([^?]+\?)", body)
    if question_sentences:
        return clean_line(question_sentences[-1])
    return "Answer the question shown in the image."


def crop_question_image(
    page_image: Path,
    output_path: Path,
    answer_top: int | None,
    jpeg_quality: int,
) -> None:
    image = Image.open(page_image).convert("RGB")
    scale_x = image.width / PAGE_XML_WIDTH
    scale_y = image.height / PAGE_XML_HEIGHT
    top = max(0, int(QUESTION_TOP * scale_y))
    bottom_xml = answer_top - 12 if answer_top else 720
    bottom = min(image.height, max(top + 120, int(bottom_xml * scale_y)))
    # Light side crop removes page margins while preserving all question text.
    left = int(10 * scale_x)
    right = int(875 * scale_x)
    cropped = image.crop((left, top, right, bottom))
    cropped.save(output_path, "JPEG", quality=jpeg_quality, optimize=True)


def add_bold_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run_obj = paragraph.add_run(text)
    run_obj.bold = True


def build_docx(
    pages: list[str],
    rendered_pages: list[Path],
    answer_tops: dict[int, int],
    out_docx: Path,
    out_txt: Path,
    work_dir: Path,
    jpeg_quality: int,
    default_difficulty: str | None,
) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    mapped_lines: list[str] = []
    question_count = 0
    short_answer_count = 0
    mc_count = 0

    for idx, page_text in enumerate(pages, start=1):
        if idx > len(rendered_pages):
            break
        if "Question ID" not in page_text:
            continue

        qid = question_id(page_text, idx)
        question_part, answer_part = split_question_answer(page_text, qid)
        options = extract_options(question_part)
        correct = extract_correct_answer(answer_part)
        prompt = extract_prompt(question_part)
        answers = accepted_answers(answer_part, correct)
        rationale = extract_rationale(answer_part)
        difficulty = extract_difficulty(answer_part) or default_difficulty
        category = extract_category(answer_part)
        is_multiple_choice = bool(options) and bool(correct) and re.fullmatch(r"[A-D]", correct.strip(), re.I)

        question_count += 1
        image_path = work_dir / f"question-{question_count:03d}-{qid}.jpg"
        crop_question_image(rendered_pages[idx - 1], image_path, answer_tops.get(idx), jpeg_quality)

        header = f"00000_Question_{question_count}({qid})_{category}{'_SC' if is_multiple_choice else ''}"
        add_bold_paragraph(document, header)
        document.add_paragraph("00001_See the original SAT Math question image below.")
        document.add_picture(str(image_path), width=Inches(6.8))
        document.add_paragraph(f"00002_{prompt}")

        mapped_lines.extend([
            header,
            "00001_See the original SAT Math question image below.",
            f"[IMAGE: {image_path.name}]",
            f"00002_{prompt}",
        ])

        if rationale:
            document.add_paragraph(f"00004_{rationale}")
            mapped_lines.append(f"00004_{rationale}")

        if difficulty:
            document.add_paragraph(f"00005_{difficulty}")
            mapped_lines.append(f"00005_{difficulty}")

        if is_multiple_choice:
            mc_count += 1
            correct_label = correct.strip().upper()
            document.add_paragraph("00006_")
            mapped_lines.append("00006_")
            option_by_label = {label: content for label, content in options}
            for label in ["A", "B", "C", "D"]:
                content = option_by_label.get(label, f"Choice {label} (see image)")
                suffix = " T (True)" if label == correct_label else ""
                document.add_paragraph(f"{content}{suffix}")
                mapped_lines.append(f"{content}{suffix}")
        else:
            short_answer_count += 1
            if not answers and correct:
                answers = [correct]
            if not answers:
                answers = ["[accepted answer needs review]"]
            answer_line = " | ".join(answers)
            document.add_paragraph(f"00003_{answer_line}")
            mapped_lines.append(f"00003_{answer_line}")

        document.add_paragraph("==End==")
        mapped_lines.append("==End==")
        mapped_lines.append("")
        if question_count != len(rendered_pages):
            document.add_page_break()

    document.save(out_docx)
    out_txt.write_text("\n".join(mapped_lines), encoding="utf-8")
    print(f"questions={question_count} multiple_choice={mc_count} short_answer={short_answer_count}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("pdf", type=Path)
    parser.add_argument("--out-docx", type=Path, required=True)
    parser.add_argument("--out-txt", type=Path, required=True)
    parser.add_argument("--dpi", type=int, default=100)
    parser.add_argument("--jpeg-quality", type=int, default=78)
    args = parser.parse_args()

    if not shutil.which("pdftotext") or not shutil.which("pdftoppm") or not shutil.which("pdftohtml"):
        raise SystemExit("pdftotext, pdftoppm, and pdftohtml are required")

    with tempfile.TemporaryDirectory(prefix="sat-pdf-convert-") as tmp:
        work_dir = Path(tmp)
        pages = read_pages(args.pdf)
        answer_tops = extract_answer_tops(args.pdf, work_dir)
        image_dir = work_dir / "rendered"
        image_dir.mkdir()
        rendered_pages = render_pages(args.pdf, image_dir, args.dpi)
        build_docx(
            pages=pages,
            rendered_pages=rendered_pages,
            answer_tops=answer_tops,
            out_docx=args.out_docx,
            out_txt=args.out_txt,
            work_dir=work_dir,
            jpeg_quality=args.jpeg_quality,
            default_difficulty=infer_default_difficulty(args.pdf),
        )


if __name__ == "__main__":
    main()
